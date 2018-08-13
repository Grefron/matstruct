import math
import itertools

import numpy as np

from scipy import interpolate as interpolate
from scipy.spatial.qhull import ConvexHull

from matplotlib import pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure

from PyQt5 import QtCore
from PyQt5 import QtWidgets

import pint

### https://stackoverflow.com/questions/24805671/how-to-use-python-docx-to-replace-text-in-a-word-document-and-save

import re

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def docx_replace_dict(doc, dictionary):
    for word, replacement in dictionary.items():
        word_re=re.compile(word)
        docx_replace_regex(doc, word_re , replacement)

###



ureg = pint.UnitRegistry()


class MultiTabNavTool(NavigationToolbar):
    def __init__(self, canvases, tabs, parent=None):
        self.canvases = canvases
        self.tabs = tabs

        NavigationToolbar.__init__(self, canvases[0], parent)

    @property
    def canvas(self):
        return self.canvases[self.tabs.currentIndex()]

    @canvas.setter
    def canvas(self, canvas):
        self._canvas = canvas


class AshbyConfigurator(QtWidgets.QWidget):
    def __init__(self, ashby):
        super().__init__()
        self.ashby = ashby


# x_prop, y_prop, group_by=None, filter_by=None, envelope=True, logscale=True, colors=None,
#              title=None, width=None, height=None, color_by=None

class MplMultiTab(QtWidgets.QMainWindow):
    def __init__(self, parent=None, figures=None, labels=None):
        QtWidgets.QMainWindow.__init__(self)

        self.main_frame = QtWidgets.QWidget()
        self.tabWidget = QtWidgets.QTabWidget(self.main_frame)
        self.create_tabs(figures, labels)

        # Create the navigation toolbar, tied to the canvas
        self.mpl_toolbar = MultiTabNavTool(self.canvases, self.tabWidget, self.main_frame)

        self.vbox = vbox = QtWidgets.QVBoxLayout()
        vbox.addWidget(self.mpl_toolbar)
        vbox.addWidget(self.tabWidget)

        self.main_frame.setLayout(vbox)
        self.setCentralWidget(self.main_frame)

    def create_tabs(self, figures, labels):
        if labels is None:
            labels = []
        figures = [
            Figure()] if figures is None else figures  # initialise with empty figure in first tab if no figures provided
        self.canvases = [self.add_tab(fig, lbl)
                         for (fig, lbl) in itertools.zip_longest(figures, labels)]

    def add_tab(self, fig=None, name=None):
        '''dynamically add tabs with embedded matplotlib canvas with this function.'''

        # Create the mpl Figure and FigCanvas objects.
        if fig is None:
            fig = Figure()
            ax = fig.add_subplot(111)

        canvas = fig.canvas if fig.canvas else FigureCanvas(fig)
        canvas.get_default_filename = lambda: "{}.png".format(name)
        canvas.setParent(self.tabWidget)
        canvas.setFocusPolicy(QtCore.Qt.ClickFocus)

        # self.tabs.append( tab )
        name = 'Tab %i' % (self.tabWidget.count() + 1) if name is None else name
        self.tabWidget.addTab(canvas, name)

        return canvas


class AttributeFilter:
    def __init__(self, equal=None, minimum=None, maximum=None):
        self.equal = equal
        self.minimum = minimum
        self.maximum = maximum

    def passes(self, obj):
        if self.equal is not None:
            for attr, val in self.equal.items():
                if isinstance(val, str):  # single criterion
                    val = [val]
                if not getattr(obj, attr) in val:
                    return False
        if self.minimum is not None:
            for attr, min_val in self.minimum.items():
                if not getattr(obj, attr) > min_val:
                    return False
        if self.maximum is not None:
            for attr, max_val in self.maximum.items():
                if not getattr(obj, attr) < max_val:
                    return False
        return True


class Material:
    def __init__(self, name, supplier=None, base_material=None, notes="",
                 density=0,
                 tensile_strength=0,
                 compressive_strength=0, compressive_modulus=0,
                 youngs_modulus=0, shear_strength=0, shear_modulus=0,
                 color='b',
                 category=None,
                 price=0):
        self.name = name
        self.supplier = supplier
        self.base_material = base_material
        self.notes = notes
        self.density = density
        self.tensile_strength = tensile_strength
        self.compressive_strength = compressive_strength
        self.compressive_modulus = compressive_modulus
        self.youngs_modulus = youngs_modulus
        self.shear_strength = shear_strength
        self.shear_modulus = shear_modulus
        self.color = color
        self.category = category
        self.price = price


class Ashby:
    # class attribute name: (name str, stored unit, display unit)
    props = {'density': ("Density", "kg/m^3", "kg/m^3"),  # kg/m3
             'tensile_strength': ("Tensile strength", "MPa", "MPa"),  # MPa
             'compressive_strength': ("Compressive strength", "MPa", "MPa"),  # MPa
             'compressive_modulus': ("Compressive modulus", "MPa", "GPa"),  # GPa
             'youngs_modulus': ("Young's modulus", "MPa", "GPa"),  # GPa
             'shear_strength': ("Shear strength", "MPa", "MPa"),  # MPa
             'shear_modulus': ("Shear modulus", "MPa", "GPa"),
             'base_material': ("Chemical base", "", ""),
             'name': ("Name", "", ""),
             'supplier': ("Supplier", "", ""),
             'price': ("Price", "1/kg", "1/kg")}  # GPa

    def __init__(self, materials, x_prop=None, y_prop=None, filter=None, group_by=None, title=None):
        self.materials = materials
        self.x_prop = x_prop
        self.y_prop = y_prop
        self.group_by = group_by
        self.title = title
        if filter is None:
            filter = AttributeFilter()
        self.filter = filter

    @property
    def materials(self):
        try:
            return self._materials.select()  # works with Ashby object
        except AttributeError:
            return self._materials

    @materials.setter
    def materials(self, value):
        self._materials = value

    def get_options(self, attr):
        return list(set(getattr(m, attr) for m in self.materials))

    def select(self):
        if self.filter is None:
            return self.materials
        return [m for m in self.materials if self.filter.passes(m)]

    def table(self, cols):
        data = self.selected_material_data()
        header = []
        for col in cols:
            display_name, unit, display_unit = Ashby.props[col]
            if display_unit == "":
                header.append(display_name)
            else:
                header.append("{} [{}]".format(display_name, display_unit))
        rows = [header]
        for _, dataset in data.items():
            for d in dataset:
                row_data = []
                for col in cols:
                    try:
                        row_data.append(self.factor(col) * getattr(d, col))
                    except (KeyError, TypeError):
                        row_data.append(getattr(d, col))
                rows.append(row_data)
        return rows

    def factor(self, prop):
        display_name, unit, display_unit = Ashby.props[prop]
        return ureg.convert(1, unit, display_unit)

    def plot(self, width=10, height=8, envelope=True, logscale=True, color_by=None, colors=None):
        figure = plt.figure(figsize=(width, height))
        ax = figure.add_subplot(111)
        ax.clear()

        def to_label(prop, units=False):
            display_name, unit, display_unit = Ashby.props[prop]
            if units:
                return "{} [{}]".format(display_name, display_unit)
            else:
                return display_name

        data = self.selected_material_data()

        ax.set_xlabel(to_label(self.x_prop, units=True))
        ax.set_ylabel(to_label(self.y_prop, units=True))
        if self.title is None:
            if self.group_by is not None:
                title = "{} vs {} by {}".format(to_label(self.y_prop, units=False), to_label(self.x_prop, units=False),
                                                self.group_by.replace("_", " "))
            else:
                title = "{} vs {}".format(to_label(self.y_prop, units=False), to_label(self.x_prop, units=False))
        else:
            title = self.title
        ax.set_title(title)
        ax.grid(True, which="both", ls="-")

        f_x = self.factor(self.x_prop)
        f_y = self.factor(self.y_prop)

        z_order = 1
        markers = ['o', 'x', '+', '*', 's']
        color_marker_index = {}
        for label, dataset in data.items():
            x_arr = [f_x * float(getattr(m, self.x_prop)) for m in dataset]
            y_arr = [f_y * float(getattr(m, self.y_prop)) for m in dataset]
            x = np.array(x_arr, dtype=float)
            y = np.array(y_arr, dtype=float)

            if colors is None:
                c = 'k'
            else:
                if color_by is None:
                    color_label = label
                else:
                    color_label = getattr(dataset[0], color_by)
                if color_label in colors:
                    c = colors[color_label]
                else:
                    c = 'k'
            if c in color_marker_index:
                color_marker_index[c] = (color_marker_index[c] + 1) % len(markers)
            else:
                color_marker_index[c] = 0
            marker = markers[color_marker_index[c]]
            ax.plot(x, y, marker=marker, color=c, linestyle="None", zorder=z_order, label=label)

            if envelope:
                contour_padding = 0.05
                f1 = 1 - contour_padding * 0.71
                f2 = 1 + contour_padding * 0.71
                f3 = 1 - contour_padding
                f4 = 1 + contour_padding
                xc = np.array([f1 * xi for xi in x_arr] +
                              [f2 * xi for xi in x_arr] +
                              [f1 * xi for xi in x_arr] +
                              [f2 * xi for xi in x_arr] +
                              [f3 * xi for xi in x_arr] +
                              [xi for xi in x_arr] +
                              [f4 * xi for xi in x_arr] +
                              [xi for xi in x_arr])
                yc = np.array([f1 * yi for yi in y_arr] +
                              [f2 * yi for yi in y_arr] +
                              [f2 * yi for yi in y_arr] +
                              [f1 * yi for yi in y_arr] +
                              [yi for yi in y_arr] +
                              [f3 * yi for yi in y_arr] +
                              [yi for yi in y_arr] +
                              [f4 * yi for yi in y_arr])
                # ax.scatter(xc, yc, marker="x", color='k')
                xi, yi = contour(xc, yc)
                ax.fill(xi, yi, color=c, alpha=0.3, zorder=z_order)
                ax.plot(xi, yi, linestyle='-', color=c, zorder=z_order)

            z_order += 1

        if logscale:
            ax.set_yscale("log")
            ax.set_xscale("log")
        ax.legend(loc=2)

        figure.tight_layout()
        return figure

    def selected_material_data(self):
        selection = self.select()
        if self.group_by is None:
            data = {'All': selection}
        else:
            data = {}
            for m in selection:
                label = getattr(m, self.group_by)
                if label in data:
                    data[label].append(m)
                else:
                    data[label] = [m]
        return data


def contour(x, y):
    # convex = convex_hull(x, y, offset=1.1)
    points = np.ndarray((len(x), 2))
    points[:, 0] = np.array(x)
    points[:, 1] = np.array(y)
    hull = ConvexHull(points)
    convex = []
    for v in hull.vertices:
        convex.append(points[v].tolist())

    # https://stackoverflow.com/questions/47948453/scipy-interpolate-splprep-error-invalid-inputs
    # Fitpack has a fit if it two consecutive inputs are identical. The error happens deep enough that it depends on how the libraries were compiled and linked, hence the assortment of errors.
    prev_xc, prev_yc = convex[-1][0], convex[-1][1]
    x_conv, y_conv = [], []
    for c in convex:
        xc, yc = c[0], c[1]
        if not ((xc == prev_xc) and (yc == prev_yc)):
            x_conv.append(xc)
            y_conv.append(yc)
        prev_xc, prev_yc = xc, yc

    # x_conv, y_conv = [p[0] for p in convex], [p[1] for p in convex]

    x_conv.append(x_conv[0])
    y_conv.append(y_conv[0])

    # taken from https://stackoverflow.com/questions/33962717/interpolating-a-closed-curve-using-scipy
    # fit splines to x=f(u) and y=g(u), treating both as periodic. also note that s=0
    # is needed in order to force the spline fit to pass through all the input points.

    base = 10
    log_x = [math.log(xi, base) for xi in x_conv]
    log_y = [math.log(yi, base) for yi in y_conv]

    # tck, u = interpolate.splprep([np.array(x_conv), np.array(y_conv)], s=0, per=True, k=3)
    tck, u = interpolate.splprep([np.array(log_x), np.array(log_y)], s=0, per=True, k=1)

    # evaluate the spline fits for n evenly spaced distance values
    xi, yi = interpolate.splev(np.linspace(0, 1, 1000), tck)

    xi = [base ** xii for xii in xi]
    yi = [base ** yii for yii in yi]

    return xi, yi


def import_materials(fname, mapping, tab):
    import openpyxl

    sheet = openpyxl.load_workbook(fname, data_only=True)[tab]

    def import_row(row, mapping):
        attrs = {}
        for attr, col in mapping.items():
            attrs[attr] = row[col].value
        return Material(**attrs)

    materials = []
    for i, row in enumerate(sheet.rows):
        if i > 0:
            materials.append(import_row(row, mapping))

    return materials


def docx_table(doc, table, table_cols):
    doc_table = doc.add_table(rows=0, cols=len(table_cols))
    # doc_table.style = 'Light List Accent 1'
    for row in table:
        row_cells = doc_table.add_row().cells
        for j, cell in enumerate(row):
            try:
                row_cells[j].text = "{:.3f}".format(cell)
            except:
                row_cells[j].text = str(cell)


if __name__ == "__main__":
    import sys
    import docx

    text = """
Remarks:

Not all mechanical properties could be found in all datasheets, in these cases the missing values have been estimated, as indicated in the “Comment” column. If no processing or use temperature is known, this value was left empty.
If multiple values for a certain property are given (e.g. if multiple test standards have been used or if values for different temperatures are given), the lowest (most conservative) value was included in the table.
Three sets of graphs are provided:

1.	Data from datasheets only, plotting mechanical properties vs density and strength vs modulus; In all cases, results are grouped per “base material” (e.g. PUR or PEI) and per “category” (e.g. Evonik Rohacell or Gurit PVC)
2.	As 1) but including relevant materials from the CES Selector database, showing the placement of the provided materials wrt a broader spectrum of materials;
3.	Data from the CES selector database, showing mechanical properties vs price data;
4.	Data from the CES selector database, showing compressive strength vs max. use temperature, with a selection for the relevant process conditions in RTM (T>100C, p>2,5MPa)
"""

    gui = False
    docx_file = True

    mapping_specific = {'supplier': 0,  # col, unit, display_unit, factor
               'name': 1,
               'category': 2,
               'base_material': 3,
               'notes': 14,
               'density': 8,  # kg/m3
               'tensile_strength': 9,  # MPa
               'compressive_strength': 10,  # MPa
               'compressive_modulus': 11,  # GPa
               'youngs_modulus': 11,  # GPa
               'shear_strength': 12,  # MPa
               'shear_modulus': 13}  # GPa

    mapping_generic = {'supplier': 1,
                       'name': 0,
                       'category': 2,
                       'base_material': 3,
                       'price': 4,
                       'density': 5,
                       'youngs_modulus': 7,
                       'compressive_modulus': 7,
                       'shear_strength': 8,
                       'tensile_strength': 9,
                       'compressive_strength': 10,
                       'shear_modulus': 12}

    colors = {'PUR': 'gold',
              'PVC': 'green',
              'PEI': 'blue',
              'PET': 'orange',
              'PMI': 'red',
              'SAN': 'purple',
              'Balsa': 'brown'}

    a = Ashby(import_materials('material_data.xlsx', mapping_specific, tab='materials'))
    b = Ashby(import_materials('material_data.xlsx', mapping_generic, tab='ces'))

    base_materials = a.get_options('base_material')
    categories = a.get_options('category')
    print(base_materials)
    print(categories)

    figs = [{'x_prop': 'density', 'y_prop': 'youngs_modulus', 'group_by': 'base_material', 'equal': {'base_material': ['SAN', 'PET', 'PMI']}, 'minimum': {'youngs_modulus': 100}},
            {'x_prop': 'density', 'y_prop': 'youngs_modulus', 'group_by': 'category', 'color_by': 'base_material'}]

    figures = []
    for args in figs:
        a.x_prop, a.y_prop = args['x_prop'], args['y_prop']
        a.group_by = args.get('group_by', None)
        a.filter.equal = args.get('equal', None)
        a.filter.minimum = args.get('minimum', None)
        a.filter.maximum = args.get('maximum', None)
        # a.filter_by = args.get('filter_by', None)
        figures.append(a.plot(colors=colors, color_by=args.get('color_by', None)))


    # figures.append(a.plot('density', 'shear_modulus', group_by='base_material', colors=colors, width=11, height=7.5))  # filter_by={'base_material': 'SAN'}
    # figures.append(a.plot('density', 'shear_modulus', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))  # filter_by={'base_material': 'SAN'}
    #
    # figures.append(a.plot('density', 'tensile_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('density', 'tensile_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))
    #
    # figures.append(a.plot('density', 'compressive_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('density', 'compressive_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))
    #
    # figures.append(a.plot('density', 'shear_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('density', 'shear_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))
    #
    # figures.append(a.plot('youngs_modulus', 'tensile_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('youngs_modulus', 'tensile_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))
    #
    # figures.append(a.plot('youngs_modulus', 'compressive_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('youngs_modulus', 'compressive_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))
    #
    # figures.append(a.plot('youngs_modulus', 'shear_strength', group_by='base_material', colors=colors, width=11, height=7.5))
    # figures.append(a.plot('youngs_modulus', 'shear_strength', group_by='category', color_by='base_material', colors=colors, width=11, height=7.5))


    def get_label(fig):
        return str(fig.axes[0].title._text)

    if gui:
        app = QtWidgets.QApplication(sys.argv)
        ui = MplMultiTab(figures=figures, labels=[get_label(fig) for fig in figures])
        ui.show()
        app.exec_()

    if docx_file:
        import io
        from PIL import Image

        table_cols = ['name', 'supplier', 'base_material', 'density',
                      'youngs_modulus', 'compressive_modulus', 'shear_modulus',
                      'tensile_strength', 'compressive_strength', 'shear_strength']

        doc = docx.Document(r"docs\template.docx")

        docx_replace_dict(doc, {"Field:DocType": "Material selection report",
                                "Field:Title": "",
                                "Field:CustomerShortName": "",
                                "Field:CustomerReference": "",
                                "Field:SirrisReference": "",
                                "Field:DocumentNumber": "",
                                "Field:DocumentVersion": "",
                                "Field:DocumentStatus": "",
                                "Field:Author": ""})

        # section = doc.sections[0]
        # header = section.header
        # header.add_picture(r"docs\sirris.png", height=docx.shared.Cm(2))

        doc.add_paragraph(text)

        for fig in figures:
            fname = r"pics\{}.png".format(get_label(fig))
            f = io.BytesIO()
            # figures[0].savefig(r"pics\test_fig.png", dpi=150, orientation='landscape')
            fig.savefig(f, dpi=150)

            im = Image.open(f)
            im.rotate(90, expand=1).save(fname)

            doc.add_picture(fname, width=docx.shared.Cm(15),
                            height=docx.shared.Cm(22))

        a.filter = None
        a.group_by = 'base_material'
        docx_table(doc, a.table(cols=table_cols), table_cols)
        docx_table(doc, b.table(cols=table_cols), table_cols)

        doc.save(r"docs\report.docx")